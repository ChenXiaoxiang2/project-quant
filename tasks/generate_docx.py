from docx import Document
import pandas as pd
import numpy as np
from datetime import datetime
import os

def create_market_report_docx(output_path, top_gainers, top_100):
    doc = Document()
    date_str = datetime.now().strftime("%Y-%m-%d")
    doc.add_heading(f'每日股市复盘报告 - {date_str}', 0)
    
    doc.add_heading('今日涨幅前十', level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '代码'
    hdr_cells[1].text = '名称'
    hdr_cells[2].text = '涨幅'
    hdr_cells[3].text = '收盘'
    
    for _, row in top_gainers.iterrows():
        row_cells = table.add_row().cells
        row_cells[0].text = str(row['ts_code'])
        row_cells[1].text = str(row['name'])
        row_cells[2].text = f"{row['pct_chg']:.2%}"
        row_cells[3].text = f"{row['close']:.2f}"
        
    doc.add_heading('潜力个股推荐 (Top 100)', level=1)
    table2 = doc.add_table(rows=1, cols=4)
    table2.style = 'Table Grid'
    hdr_cells2 = table2.rows[0].cells
    hdr_cells2[0].text = '代码'
    hdr_cells2[1].text = '名称'
    hdr_cells2[2].text = '板块'
    hdr_cells2[3].text = '收盘'
    
    for _, row in top_100.iterrows():
        row_cells = table2.add_row().cells
        row_cells[0].text = str(row['ts_code'])
        row_cells[1].text = str(row['name'])
        row_cells[2].text = str(row['sector'])
        row_cells[3].text = f"{row['close']:.2f}"

    doc.save(output_path)
    print(f"Word报告已生成: {output_path}")

if __name__ == '__main__':
    # 模拟数据生成
    data = {
        'ts_code': [f'600{i:03}' for i in range(200)],
        'name': [f'股票{i}' for i in range(200)],
        'sector': np.random.choice(['新能源', '人工智能', '半导体', '医药'], 200),
        'pct_chg': np.random.uniform(-0.1, 0.1, 200),
        'close': np.random.uniform(10, 100, 200),
        'score': np.random.rand(200)
    }
    df = pd.DataFrame(data)
    create_market_report_docx(
        './reports/daily_report_20260321.docx', 
        df.nlargest(10, 'pct_chg'), 
        df.nlargest(100, 'score')
    )
